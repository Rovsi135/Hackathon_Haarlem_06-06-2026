"""
generate_bites.py
-----------------
Generates pre-bite and post-bite documents as .docx files using python-docx.

Input:  validated_input.json   (intake from the intake agent)
        slides_init_text.json  (generated slides content)
Output: pre_bite.docx
        post_bite.docx

Run this after slides_generator.py has completed.
Install: pip install python-docx
"""

import json
import os
import sys
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.append("../pptx_generator")
from style_config import COLORS, FONTS

load_dotenv()

client = OpenAI(
    api_key=os.environ["OPENROUTER_API_KEY"],
    base_url="https://openrouter.ai/api/v1",
)

# ── Colors from style_config (strip # for python-docx XML) ───────────────────

def hex_str(color_key: str) -> str:
    """Look up a color by name and return it as a bare hex string (no #)."""
    return COLORS[color_key].lstrip("#")

C_PURPLE    = hex_str("dark_purple")   # 1A0040 — header background
C_BODY_BLUE = hex_str("primary_dark")  # 0D006A — section headings
C_RED       = hex_str("rose_red")      # EF4453 — pre-bite tag
C_TEAL      = hex_str("teal")          # 00B0F0 — post-bite tag
C_LAVENDER  = hex_str("bg_lavender")   # EDE9FF — highlight box
C_BORDER    = hex_str("lavender_tint") # BCB3FF — heading underline
C_GREY      = hex_str("dark_grey")     # 262626 — muted text
C_WHITE     = "FFFFFF"                 # Pure white — not in palette, needed for text on dark bg

# ── Fonts from style_config ───────────────────────────────────────────────────

FONT_BODY  = FONTS["primary"]   # Space Grotesk (falls back gracefully in Word if not installed)
FONT_ALT   = FONTS["fallback"]  # Calibri

def rgb(hex_str: str) -> RGBColor:
    """Convert bare hex string to RGBColor."""
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

# ── Load inputs ───────────────────────────────────────────────────────────────

with open("validated_input.json", "r", encoding="utf-8") as f:
    intake = json.load(f)

with open("slides_init_text.json", "r", encoding="utf-8") as f:
    slides_data = json.load(f)

slide_summary = "\n".join(
    f"[{s['block'].upper()}] {s['title']}: {', '.join(s['bullets'][:3])}"
    for s in slides_data["slides"]
)

# ── LLM call ──────────────────────────────────────────────────────────────────

BITES_SYSTEM_PROMPT = """
You are a professional training designer for Maverx, a Dutch consultancy.
You write clear, practical pre-session and post-session participant documents.

A pre-bite prepares participants BEFORE the training. It must contain:
- A short context paragraph (2-3 sentences) explaining why this training matters
- A reflection question to think about before attending
- A concrete preparation task (install software, read a short article, or answer a prompt)
- A recommended resource (article, video, or tool) with a brief description

A post-bite follows up AFTER the training. It must contain:
- A short recap paragraph (2-3 sentences) of what was covered
- 3 reflection questions tied to the learning objective
- A practical assignment to apply the learning within 1 week
- 2 further reading items for deepening knowledge

RULES:
- Write in the same language register as the audience (professional but approachable)
- Be specific — reference the actual training topic, not generic placeholders
- The assignment must be concrete and completable, not vague
- Output ONLY valid JSON, no markdown, no explanation, no preamble

OUTPUT FORMAT:
{
  "pre_bite": {
    "title": "...",
    "context": "...",
    "reflection_question": "...",
    "preparation_task": "...",
    "optional_resource": {
      "title": "...",
      "description": "...",
      "url_or_note": "..."
    }
  },
  "post_bite": {
    "title": "...",
    "recap": "...",
    "reflection_questions": ["...", "...", "..."],
    "assignment": {
      "title": "...",
      "description": "...",
      "deadline": "Within 1 week after the session"
    },
    "further_reading": [
      { "title": "...", "description": "..." },
      { "title": "...", "description": "..." }
    ]
  }
}
"""

BITES_USER_PROMPT = f"""
Generate the pre-bite and post-bite documents for this training:

Topic: {intake['topic']}
Audience: {intake['audience']}
Knowledge level: {intake['knowledge_level']}
Duration: {intake['duration_hours']} hours
Learning objective: {intake['learning_objective']}

Slide content summary:
{slide_summary}
"""

print("[1/3] Generating pre-bite and post-bite content...")

response = client.chat.completions.create(
    model="anthropic/claude-sonnet-4-6",
    max_tokens=4000,
    messages=[
        {"role": "system", "content": BITES_SYSTEM_PROMPT},
        {"role": "user", "content": BITES_USER_PROMPT},
    ]
)

raw = response.choices[0].message.content
if raw.strip().startswith("```"):
    raw = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()

bites = json.loads(raw)
pre   = bites["pre_bite"]
post  = bites["post_bite"]
print("      Content generated.")

# ── python-docx helpers ───────────────────────────────────────────────────────

def shade_paragraph(p, hex_color: str):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_header_bar(doc, title, subtitle, bar_color, tag_text, tag_color):
    # Tag (e.g. PRE-BITE), right-aligned on colored background
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = tag_p.add_run(tag_text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(C_WHITE)
    r.font.name = FONT_BODY
    shade_paragraph(tag_p, tag_color)

    # Main title on dark purple background
    title_p = doc.add_paragraph()
    r2 = title_p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(FONTS["subtitle"])
    r2.font.color.rgb = rgb(C_WHITE)
    r2.font.name = FONT_BODY
    shade_paragraph(title_p, bar_color)
    pPr = title_p._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "120")
    sp.set(qn("w:after"),  "120")
    pPr.insert(0, sp)

    # Subtitle
    sub_p = doc.add_paragraph()
    r3 = sub_p.add_run(subtitle)
    r3.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = rgb(C_GREY)
    r3.font.name = FONT_BODY
    sub_p.paragraph_format.space_after = Pt(16)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = rgb(C_BODY_BLUE)
    r.font.name = FONT_BODY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), C_BORDER)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(FONTS["body"])
    r.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(space_after)


def add_highlight_box(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(FONTS["body"])
    r.font.name = FONT_BODY
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)
    shade_paragraph(p, C_LAVENDER)


def add_bold_line(doc, text, space_after=2):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(FONTS["body"])
    r.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(space_after)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(FONTS["body"])
    r.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(4)


# ── Build documents ───────────────────────────────────────────────────────────

print("[2/3] Building .docx files...")


def build_pre_bite(data, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    add_header_bar(doc,
        title     = data["title"],
        subtitle  = f"Preparation Guide · {intake['topic']}",
        bar_color = C_PURPLE,
        tag_text  = "PRE-BITE",
        tag_color = C_RED,
    )

    add_section_heading(doc, "Why This Training?")
    add_body(doc, data["context"])

    add_section_heading(doc, "Before You Attend — Reflect on This")
    add_highlight_box(doc, data["reflection_question"])

    add_section_heading(doc, "Your Preparation Task")
    add_body(doc, data["preparation_task"])

    add_section_heading(doc, "Recommended Resource (Optional)")
    add_bold_line(doc, data["optional_resource"]["title"])
    add_body(doc, data["optional_resource"]["description"], space_after=2)
    p = doc.add_paragraph()
    r = p.add_run(data["optional_resource"]["url_or_note"])
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(C_GREY)
    r.font.name = FONT_BODY

    doc.save(output_path)


def build_post_bite(data, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    add_header_bar(doc,
        title     = data["title"],
        subtitle  = f"Follow-Up Guide · {intake['topic']}",
        bar_color = C_PURPLE,
        tag_text  = "POST-BITE",
        tag_color = C_TEAL,
    )

    add_section_heading(doc, "What You Covered Today")
    add_body(doc, data["recap"])

    add_section_heading(doc, "Reflect on Your Learning")
    for q in data["reflection_questions"]:
        add_numbered(doc, q)

    add_section_heading(doc, "Your Assignment")
    add_bold_line(doc, data["assignment"]["title"])
    add_body(doc, data["assignment"]["description"], space_after=4)
    p = doc.add_paragraph()
    r1 = p.add_run("Deadline: ")
    r1.bold = True
    r1.font.size = Pt(FONTS["body"])
    r1.font.name = FONT_BODY
    r2 = p.add_run(data["assignment"]["deadline"])
    r2.font.size = Pt(FONTS["body"])
    r2.font.name = FONT_BODY
    p.paragraph_format.space_after = Pt(8)

    add_section_heading(doc, "Further Reading")
    for item in data.get("further_reading", []):
        add_bold_line(doc, item["title"])
        add_body(doc, item["description"], space_after=8)

    doc.save(output_path)


build_pre_bite(pre,   "pre_bite.docx")
build_post_bite(post, "post_bite.docx")

print("[3/3] Done.")
print("  ✓ pre_bite.docx")
print("  ✓ post_bite.docx")